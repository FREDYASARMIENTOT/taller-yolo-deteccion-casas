from docx import Document
from docx.shared import Pt
import argparse
import os


def create_docx(readme_path, output_path, professor, students):
    if not os.path.exists(readme_path):
        raise FileNotFoundError(f"No se encontró {readme_path}")

    with open(readme_path, "r", encoding="utf-8") as f:
        readme = f.read()

    doc = Document()

    # Portada
    doc.add_heading('Taller YOLOv8 - Detección de Casas', level=0)
    p = doc.add_paragraph()
    p.add_run(f'Profesor: {professor}\n').bold = True
    p.add_run('Grupo de estudiantes:\n').bold = True
    for s in students:
        p.add_run(f'- {s}\n')

    doc.add_page_break()

    # Contenido del README
    doc.add_heading('Documentación del proyecto', level=1)

    # Convertir el README en párrafos simples (mantener saltos de línea dobles como separadores)
    blocks = readme.split('\n\n')
    for block in blocks:
        lines = block.strip().split('\n')
        for line in lines:
            # Mantener viñetas y encabezados tal cual
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(11)
        doc.add_paragraph()

    # Guardar documento
    doc.save(output_path)
    print(f'Documento guardado en: {output_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generar archivo Word (.docx) con la documentación del taller')
    parser.add_argument('--readme', type=str, default='README.md', help='Ruta al README.md')
    parser.add_argument('--output', type=str, default='Taller_YOLOv8_Deteccion_Casas.docx', help='Nombre del archivo .docx de salida')
    parser.add_argument('--professor', type=str, default='William Davis Prada', help='Nombre del profesor')
    parser.add_argument('--students', type=str, nargs='+', default=[
        'Jorge Enrique Bravo Rojas',
        'Fredy Alejandro Sarmiento Torres',
        'Manuel Alonso Caro Ospina'
    ], help='Lista de estudiantes')

    args = parser.parse_args()
    create_docx(args.readme, args.output, args.professor, args.students)
